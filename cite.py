import re
from docx import Document
from docx.shared import Pt, Inches, Length
import os

def create_bibfile(bib = f'{os.path.expanduser("~")}/workspace/bibtex/library.bib', bib_r = 'r_references.bib'):
    '''Merges two bibfiles and creates new bibfile in cwd'''
    with open(bib_r, "r") as f:
        packages = f.read()
    with open(bib, "r") as f:
        library = f.read()
    with open("references.bib", "w+") as bib_all:
        bib_all.write(packages + library)
    
    

def bib_modify(filename):

    '''
    Takes filename of a .bib file as a string, scans the file for all titles and applies correct capitalization to the titles.
    '''
    #collect indices of titles in bibfile
    indices = []
    with open(filename, 'rb') as f:
        content = f.readlines()
        decoded = [i.decode() for i in content]
        for i in range(len(decoded)):
            if decoded[i][0:5] == 'title':
                indices.append(i)
    #isolate tilte string and apply rules
    pattern = re.compile(r'({{)(.*)(}})')
    newl =[]
    for i in indices:
        cap_ind = []
        old = pattern.findall(decoded[i])[0][1].lower().capitalize()
        l = old.split()
        for j in range(len(l)):
            if l[j][-1] in ':?!.':
                cap_ind.append(j+1)
        for e in cap_ind:
            if len(l)>e:
                l[e] = l[e].capitalize()
        new = 'title = {{' + ' '.join(l) + '}}, \n'
        newl.append(new)
    #replace corrected titles
    for i in range(len(indices)):
        decoded[indices[i]] =  newl[i]
    #write to bibfile
    encoded = [i.encode() for i in decoded]
    with open(filename, 'wb') as f:
        for i in encoded:
            f.write(i)


def hang_ind(filename, doublespace=True):
    '''Takes a docx filename in cwd as string, searches for a paragraph that equals 'References'
    in that document and adds hanging ident. Default: Adds also double spacing to all following paragraphs. Saves output in new file. Sourcefile unchanged.'''
    
    #load document
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list
    for p in range(n):
        if doc.paragraphs[p].text == 'References':
            start = p+1
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, n):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        if doublespace:
            doc.paragraphs[p].paragraph_format.line_spacing = 2.0
    name_split = filename.split('.')
    doc.save(f'{name_split[0]}_apa.{name_split[-1]}')



def rb_hang_ind(filename, doublespace=True):
    '''Takes a docx filename in cwd as string, searches for a paragraph that equals 'References'
    in that document and adds hanging ident. Default: Adds also double spacing to all following paragraphs. 
    Saves output in new file. Sourcefile unchanged.'''
    #load document
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list
    for p in range(n):
        if doc.paragraphs[p].text == '2d. Literature eferences':
            start = p+1
        if doc.paragraphs[p].text == '2e. Time Plan':
            end = p
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, end):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        if doublespace:
            doc.paragraphs[p].paragraph_format.line_spacing = 2.0

    name_split = filename.split('.')
    doc.save(f'{name_split[0]}_apa.{name_split[-1]}')






